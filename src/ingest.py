"""
MODULE 1: UNIVERSAL DOCUMENT INGESTION
----------------------------------------
Reads PDF, DOCX, Excel/CSV, and scanned (image-based) documents from /data,
converts everything into plain text chunks with rich metadata, embeds them,
and builds the FAISS index.

This REPLACES the old ingest.py (PDF-only). Same output files, more input types.

Usage:
    python src/ingest.py
"""

import os
import json
import fitz          # PyMuPDF - text PDFs
import docx           # python-docx - Word files
import pandas as pd    # Excel/CSV
import pytesseract     # OCR
from pdf2image import convert_from_path
from PIL import Image
import faiss
from sentence_transformers import SentenceTransformer

DATA_DIR = "data"
INDEX_PATH = "outputs_store/faiss_index.bin"
METADATA_PATH = "outputs_store/metadata.json"

CHUNK_SIZE = 800
CHUNK_OVERLAP = 150

# Minimum characters of extractable text per page before we assume it's a
# scanned image and fall back to OCR instead.
OCR_FALLBACK_THRESHOLD = 20


# ---------------------------------------------------------------------------
# Per-format readers. Each returns a list of {"text": ..., "page": ...} dicts
# so every format flows into the same chunking step below.
# ---------------------------------------------------------------------------

def read_pdf(filepath):
    """
    Read a PDF. If a page has almost no extractable text (a scanned page),
    fall back to OCR on that page's rendered image.
    """
    doc = fitz.open(filepath)
    pages = []
    needs_ocr_pages = []

    for page_num, page in enumerate(doc):
        text = page.get_text().strip()
        if len(text) < OCR_FALLBACK_THRESHOLD:
            needs_ocr_pages.append(page_num)
        else:
            pages.append({"page": page_num + 1, "text": text})
    doc.close()

    if needs_ocr_pages:
        print(f"    {len(needs_ocr_pages)} page(s) look scanned — running OCR...")
        try:
            images = convert_from_path(filepath)
            for page_num in needs_ocr_pages:
                if page_num < len(images):
                    ocr_text = pytesseract.image_to_string(images[page_num])
                    if ocr_text.strip():
                        pages.append({"page": page_num + 1, "text": ocr_text})
        except Exception as e:
            print(f"    OCR failed ({e}) — skipping scanned pages. "
                  f"Install poppler-utils + tesseract-ocr if this persists.")

    pages.sort(key=lambda p: p["page"])
    return pages


def read_docx(filepath):
    """Read a Word document. Word docs don't have 'pages' in the file format,
    so we treat the whole document as one logical unit and let chunking split it."""
    d = docx.Document(filepath)
    full_text = "\n".join(p.text for p in d.paragraphs if p.text.strip())

    # Also pull any tables (common in inspection/maintenance DOCX templates)
    for table in d.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text:
                full_text += "\n" + row_text

    return [{"page": 1, "text": full_text}] if full_text.strip() else []


def read_spreadsheet(filepath):
    """
    Read Excel/CSV files. Each row is converted into a natural-language
    sentence so it can be embedded and retrieved like any other text,
    e.g. 'Equipment: Pump P101 | Failure: Leakage | Date: Jan 2025'.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".csv":
        df = pd.read_csv(filepath)
    else:
        df = pd.read_excel(filepath)

    rows_as_text = []
    for _, row in df.iterrows():
        row_sentence = " | ".join(
            f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])
        )
        if row_sentence.strip():
            rows_as_text.append(row_sentence)

    # Group rows into page-sized blocks so we don't create one microscopic
    # chunk per spreadsheet row.
    pages = []
    rows_per_page = 15
    for i in range(0, len(rows_as_text), rows_per_page):
        block = "\n".join(rows_as_text[i:i + rows_per_page])
        pages.append({"page": (i // rows_per_page) + 1, "text": block})
    return pages


def read_image(filepath):
    """Pure image file (e.g. a scanned inspection form saved as .png/.jpg) via OCR."""
    text = pytesseract.image_to_string(Image.open(filepath))
    return [{"page": 1, "text": text}] if text.strip() else []


READERS = {
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".xlsx": read_spreadsheet,
    ".xls": read_spreadsheet,
    ".csv": read_spreadsheet,
    ".png": read_image,
    ".jpg": read_image,
    ".jpeg": read_image,
}


def infer_category(filename, text_sample):
    """
    Lightweight rule-based tagging so every chunk carries a rough category
    even before the LLM-based entity extraction module runs. Used for
    filtering in the dashboard and for faster, cheaper retrieval later.
    """
    name = filename.lower()
    sample = text_sample.lower()
    if any(k in name or k in sample for k in ["oisd", "peso", "factory act", "regulation", "compliance"]):
        return "Regulatory"
    if any(k in name or k in sample for k in ["maintenance", "work order", "repair"]):
        return "Maintenance"
    if any(k in name or k in sample for k in ["inspection", "audit"]):
        return "Inspection"
    if any(k in name or k in sample for k in ["safety", "permit", "sop", "procedure"]):
        return "Safety Procedure"
    if any(k in name or k in sample for k in ["incident", "near miss", "failure"]):
        return "Incident Report"
    return "General"


def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def build_index():
    print("Loading embedding model (first run downloads it, ~90MB)...")
    model = SentenceTransformer("all-MiniLM-L6-v2")

    all_chunks = []
    all_metadata = []

    files = [f for f in os.listdir(DATA_DIR)
             if os.path.splitext(f)[1].lower() in READERS]

    if not files:
        print(f"No supported files found in '{DATA_DIR}/'. "
              f"Supported: {', '.join(READERS.keys())}")
        return

    for filename in files:
        filepath = os.path.join(DATA_DIR, filename)
        ext = os.path.splitext(filename)[1].lower()
        print(f"Reading {filename} ({ext})...")

        try:
            pages = READERS[ext](filepath)
        except Exception as e:
            print(f"    Failed to read {filename}: {e}")
            continue

        for page_data in pages:
            category = infer_category(filename, page_data["text"][:300])
            chunks = chunk_text(page_data["text"])
            for i, chunk in enumerate(chunks):
                all_chunks.append(chunk)
                all_metadata.append({
                    "source": filename,
                    "page": page_data["page"],
                    "chunk_id": f"{filename}_p{page_data['page']}_c{i}",
                    "category": category,
                    "file_type": ext.replace(".", "")
                })

    if not all_chunks:
        print("No text could be extracted from any file. Check OCR setup if using scanned documents.")
        return

    print(f"Created {len(all_chunks)} chunks from {len(files)} document(s).")
    print("Generating embeddings...")
    embeddings = model.encode(all_chunks, show_progress_bar=True, convert_to_numpy=True)

    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(embeddings.astype("float32"))

    os.makedirs("outputs_store", exist_ok=True)
    faiss.write_index(index, INDEX_PATH)

    with open(METADATA_PATH, "w") as f:
        json.dump({"chunks": all_chunks, "metadata": all_metadata}, f)

    categories = {}
    for m in all_metadata:
        categories[m["category"]] = categories.get(m["category"], 0) + 1
    print("\nChunks by category:")
    for cat, count in sorted(categories.items(), key=lambda x: -x[1]):
        print(f"  {cat}: {count}")

    print(f"\nDone. Index saved to {INDEX_PATH}, metadata saved to {METADATA_PATH}")


if __name__ == "__main__":
    build_index()
