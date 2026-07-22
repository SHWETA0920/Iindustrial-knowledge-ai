"""
HACKATHON DEMO DATA GENERATOR
-----------------------------
Creates a small synthetic multi-format industrial corpus so the full product
can be demoed even without plant data.
"""

import csv
import os
from datetime import datetime

from PIL import Image, ImageDraw
import docx

DATA_DIR = "data"


def ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def write_docx(path, title, paragraphs):
    d = docx.Document()
    d.add_heading(title, level=1)
    for para in paragraphs:
        d.add_paragraph(para)
    d.save(path)


def write_csv(path, rows):
    with open(path, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_scan_image(path, lines):
    img = Image.new("RGB", (1400, 900), color="white")
    draw = ImageDraw.Draw(img)
    y = 60
    for line in lines:
        draw.text((60, y), line, fill="black")
        y += 55
    img.save(path)


def build_demo_corpus():
    ensure_dir(DATA_DIR)

    write_csv(os.path.join(DATA_DIR, "maintenance_history.csv"), [
        {"Equipment": "Pump P101", "Failure": "Seal leakage", "Date": "2019-05-12", "Cause": "Seal material degradation", "Action": "Seal replaced"},
        {"Equipment": "Pump P101", "Failure": "Seal leakage", "Date": "2022-08-17", "Cause": "Seal material degradation", "Action": "Seal replaced and alignment checked"},
        {"Equipment": "Pump P101", "Failure": "Seal leakage", "Date": "2025-02-03", "Cause": "Seal material degradation", "Action": "Seal upgraded and lubrication route reviewed"},
        {"Equipment": "Compressor C201", "Failure": "Bearing overheating", "Date": "2025-04-18", "Cause": "Poor lubrication", "Action": "Lubrication frequency increased"},
        {"Equipment": "Compressor C201", "Failure": "Bearing overheating", "Date": "2025-06-24", "Cause": "Poor lubrication", "Action": "Oil contamination removed"},
    ])

    write_docx(
        os.path.join(DATA_DIR, "plant_hot_work_sop.docx"),
        "Hot Work SOP",
        [
            "Hot work permits shall be reviewed by the area engineer before authorization.",
            "Emergency shutdown function test shall be performed every 12 months.",
            "Gas testing is mandatory before welding or cutting inside process areas.",
            "Permit issuer shall verify fire watch availability and extinguisher readiness.",
        ],
    )

    write_docx(
        os.path.join(DATA_DIR, "oisd_hot_work_guideline.docx"),
        "OISD Hot Work Guideline Excerpt",
        [
            "OISD guidance requires emergency shutdown test every 6 months for high-risk process units.",
            "Hot work permit must include gas test record, fire watch assignment, and area isolation confirmation.",
            "All permit deviations must be documented and approved by competent authority.",
        ],
    )

    write_docx(
        os.path.join(DATA_DIR, "pump_p101_manual.docx"),
        "Pump P101 Maintenance Manual",
        [
            "Pump P101 seal chamber should be inspected weekly for leakage signs.",
            "Persistent seal leakage may indicate seal material degradation or shaft misalignment.",
            "Lubrication route shall be verified after every shutdown and seal replacement.",
        ],
    )

    write_scan_image(
        os.path.join(DATA_DIR, "inspection_scan.png"),
        [
            "Inspection Record",
            "Equipment: Boiler B102",
            "Observation: Pressure drop during startup",
            "Possible cause: Valve V301 passing internally",
            "Inspector: Shift Engineer",
            f"Generated: {datetime.utcnow().date().isoformat()}",
        ],
    )

    return {
        "status": "ok",
        "files_created": sorted(os.listdir(DATA_DIR)),
        "message": "Demo corpus generated in data/",
    }


if __name__ == "__main__":
    print(build_demo_corpus())
